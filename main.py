import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from docx2pdf import convert

response = requests.get("https://ilhaad.com").text
soup_instance = BeautifulSoup(response, features="lxml")
text = soup_instance.find("aside", id="categories-2")
text = text.find_all("a")
# file = open("BaseLinks.txt", "x")
# for content in text:
#    file.write(str(content))
#    print(content)


for content in text:
    if content != text[0]:
        x = str(content)
        x = x.split('"')
        link = x[1]
        topic = x[2].split("<")[0].split(">")[1]
        formatted_text = f"Topic:-\n{topic}\nLink = {link}\n"
        print(formatted_text)
        response2 = requests.get(f"{link}").text
        soup_instance2 = BeautifulSoup(response2, features='html.parser')
        text2 = soup_instance2.find("div", id="primary")
        text2 = text2.find_all("h2")
        for sub_content in text2:
            y = str(sub_content)
            y = y.split('"')
            sub_link = y[3]
            sub_topic = y[6].split("<")[0].split(">")[1]

            sub_formatted_text = f"Sub Topic: {sub_topic}\n Sub_link = {sub_link}"
            print(sub_formatted_text)

            response3 = requests.get(f"{sub_link}").text
            soup_instance3 = BeautifulSoup(response3, features='html.parser')
            text3 = soup_instance3.find("columns")

            to_remove1 = text3.find("span")
            to_remove1.extract()

            text5 = text3.find_all("p")
            text_total = ""
            length = len(text5)
            for z in text5:
                z = z.text
                splitted = z.split("<br/>")
                for breaks in splitted:
                    text_total = text_total + "\n\n" + breaks

            text5 = str(text_total).replace('''Your email address will not be published. Required fields are marked *

Comment 

Name * 

Email * 

Website 

 Notify me of follow-up comments by email.

 Notify me of new posts by email.

 





''', "").replace('''

۔


''', "").replace('''Click to share on Facebook (Opens in new window)Click to share on Twitter (Opens in new window)Click to share on LinkedIn (Opens in new window)Click to share on Pinterest (Opens in new window)

Related''', "").replace('''Your email address will not be published. Required fields are marked *

Comment 

Name * 

Email * 

Website 

 Notify me of follow-up comments by email.

 Notify me of new posts by email.
''', "").replace('''



''', '''
''')
            text5 = text5.replace("[", "").replace(",", "").replace("'", "").replace("]", "")
            text5 = text5.replace('Your email address will not be published. Required fields are marked * Comment  '
                                  'Name *  Email *  Website   Notify me of follow-up comments by email.  Notify me of '
                                  'new posts by email.  \\n\\n  ', "").replace('n\\', "").replace('Your email address '
                                                                                                  'will not be '
                                                                                                  'published. '
                                                                                                  'Required fields '
                                                                                                  'are marked * '
                                                                                                  'Comment  Name *  '
                                                                                                  'Email *  Website   '
                                                                                                  'Notify me of '
                                                                                                  'follow-up comments '
                                                                                                  'by email.  Notify '
                                                                                                  'me of new posts by '
                                                                                                  'email.  \n', "")

            #            try:
            #                file = io.open(f"{sub_topic}.txt", "w", encoding="utf-8")
            #                file.write(f"{sub_topic}\n\n{str(text5)}")
            #                file.close()
            #            except:
            #                print("Write Error")
            print(text5)
            sub_topic.replace("/", "-").replace(":", "-").replace("<", "-").replace('"', "-").replace('n\\', "-")
            document = Document()
            title = document.add_heading(f"{sub_topic}", 0)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            para = document.add_paragraph(f"{text5}")
            para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            style = document.styles["Normal"]
            font = style.font
            font.name = "Jameel Noori Nastaleeq"
            font.size = Pt(18)

            try:
                document.save(f"{sub_topic}.docx")
            except:
                document.save(f"NameNotCorrect.docx")
            try:
                convert(f"{sub_topic}.docx", f"{sub_topic}.pdf")
            except:
                print("Converting error")
                try:
                    convert(f"{sub_topic}.docx", f"{sub_topic}_excepted.pdf")
                except:
                    print("Converting error 2")
